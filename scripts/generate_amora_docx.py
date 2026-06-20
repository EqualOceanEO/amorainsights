from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    from docx.oxml import parse_xml
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 使用XML设置背景色
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
        fill_color
    ))
    tcPr.append(shading_elm)

def create_doc():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style.font.size = Pt(11)
    
    # 标题
    title = doc.add_heading('AMORA Report Template v1.0', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('行业研究报告标准化模板')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    subtitle2 = doc.add_paragraph('适用于 AMORA 研究团队全行业报告产出')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].italic = True
    subtitle2.runs[0].font.size = Pt(10)
    subtitle2.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    # 一、模板定位
    doc.add_heading('一、模板定位', level=1)
    doc.add_paragraph('AMORA Report Template 是 AMORA 研究团队的')
    p = doc.add_paragraph('行业研究报告标准化框架')
    p.runs[0].bold = True
    doc.add_paragraph('，与企业评估框架 AMORA Score 并行使用：')
    
    # 框架对比表格
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '框架'
    hdr_cells[1].text = '用途'
    hdr_cells[2].text = '对象'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    row_cells = table.rows[1].cells
    row_cells[0].text = 'AMORA Score'
    row_cells[1].text = '企业五维评分'
    row_cells[2].text = '单个公司'
    
    row_cells = table.rows[2].cells
    row_cells[0].text = 'AMORA Report'
    row_cells[1].text = '行业研究模板'
    row_cells[2].text = '整个行业'
    
    doc.add_paragraph()
    
    # 二、五维框架定义
    doc.add_heading('二、五维框架定义', level=1)
    
    table2 = doc.add_table(rows=6, cols=5)
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    headers = ['字母', '维度', '英文', '核心问题', '关键输出']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], '2E75B6')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    data = [
        ['A', '技术先进性', 'Advancement', '技术壁垒多高？谁在领跑？卡脖子在哪？', '技术路线判断、零部件分析、专利对比'],
        ['M', '产业链生态位', 'Mapping', '上下游玩家分布？中美各自在哪占优？', '产业链图谱、生态位对比、空白机会'],
        ['O', '商业化运营', 'Operations', '应用有多深？场景成熟度？ROI？', '客户结构、落地案例、回报周期'],
        ['R', '市场容量', 'Reach', '市场多大？增速多快？天花板在哪？', '规模预测、细分拆分、增长驱动'],
        ['A', '资本价值', 'Assets', '谁更赚钱？谁更值钱？资本效率？', '财务对比、估值排名、投资建议'],
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table2.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_paragraph()
    
    # 三、报告结构规范
    doc.add_heading('三、报告结构规范', level=1)
    doc.add_heading('3.1 必含章节', level=2)
    
    sections = [
        ('封面', ['报告标题', 'AMORA 五维雷达图（可视化品牌识别）', '核心数据亮点（3-4个关键数字）', '发布日期/研究团队']),
        ('Executive Summary 执行摘要（2页）', ['核心发现（3-5条，一句话一条）', '关键数据（市场规模/增速/主要玩家）', '中美对比结论（谁在哪占优）', '战略建议（给谁看？建议做什么？）']),
        ('A. Advancement 技术先进性', ['A.1 核心发现（1段话总结本章结论）', 'A.2 技术成熟度与路线分歧', 'A.3 核心零部件分析（中美玩家+替代难度评分）', 'A.4 专利与研发投入对比', 'A.5 技术趋势判断（未来3年突破点）']),
        ('M. Mapping 产业链生态位', ['M.1 核心发现', 'M.2 产业链全景图（上游→中游→下游）', 'M.3 中美玩家生态位对比（四象限或分层）', 'M.4 供应链脆弱性评估（卡脖子清单）', 'M.5 生态位空白机会（哪里还有进入空间）']),
        ('O. Operations 商业化运营', ['O.1 核心发现', 'O.2 应用场景成熟度矩阵（量化评分）', 'O.3 客户结构深度分析', 'O.4 ROI计算与回报周期', 'O.5 落地案例对比（谁在哪真正跑通）']),
        ('R. Reach 市场容量', ['R.1 核心发现（含预测方法论说明）', 'R.2 市场规模预测方法论', 'R.3 三情景预测（保守/中性/乐观）', 'R.4 细分市场拆分', 'R.5 渗透率曲线与规模化时间线']),
        ('A. Assets 资本价值', ['A.1 核心发现', 'A.2 主要玩家财务对比', 'A.3 估值与资本效率排名', 'A.4 融资动态与估值趋势', 'A.5 盈利能力预测与投资建议']),
        ('Risk & Outlook 风险与展望', ['关键风险因素（技术/政策/市场/竞争）', '三情景展望（乐观/中性/悲观）', '时间线预测（未来1-3-5年关键节点）']),
        ('Appendix 附录', ['完整公司列表', '数据来源说明', '术语表', '参考资料']),
    ]
    
    for section_title, items in sections:
        p = doc.add_paragraph()
        p.add_run(section_title).bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # 3.2 每章必备元素
    doc.add_heading('3.2 每章必备元素', level=2)
    doc.add_paragraph('每章（A/M/O/R/A）必须包含：')
    elements = [
        ('核心发现（Key Finding）', '1段话总结本章结论，放在章首'),
        ('对比判断', '不是并列陈述，是A vs B谁赢、差多少'),
        ('数据支撑', '每个结论必须有数字支撑'),
        ('So What', '对读者意味着什么'),
    ]
    for title, desc in elements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f'：{desc}')
    
    doc.add_page_break()
    
    # 四、写作规范
    doc.add_heading('四、写作规范', level=1)
    doc.add_heading('4.1 核心发现写法', level=2)
    doc.add_paragraph('模板：').runs[0].bold = True
    p = doc.add_paragraph('[行业]处于[阶段]，核心瓶颈在[环节]而非[环节]。[国家A]在[维度]占优，[国家B]在[维度]占优。关键数据：[数字] vs [数字]，差距[X倍]。')
    p.runs[0].italic = True
    
    doc.add_paragraph('示例（人形机器人）：').runs[0].bold = True
    p = doc.add_paragraph('人形机器人处于"0.5→1.0"过渡期，核心瓶颈在触觉传感器和电池续航而非AI算法。美国在上游AI芯片/传感器占优，中国在中游制造/量产占优。关键数据：Figure融资$1.5B出货<150台，宇树融资$150M出货5500+台，资本效率差37倍。')
    p.runs[0].italic = True
    
    doc.add_heading('4.2 对比判断写法', level=2)
    table3 = doc.add_table(rows=4, cols=2)
    table3.style = 'Table Grid'
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = '错误（描述性）'
    hdr_cells[1].text = '正确（判断性）'
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '2E75B6')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    comparisons = [
        ('美国有Figure AI、Tesla、Boston Dynamics', '美国第一梯队3家，中国5家，但美国单家公司估值是中国的10倍+'),
        ('触觉传感器价格$300-500', '触觉传感器是中国产业链最大短板，2-3年内无替代可能'),
        ('市场规模预测$30B-$380B', '预测差600倍的根因是口径混乱，我们以整机出货量×均价为基准'),
    ]
    for i, (wrong, right) in enumerate(comparisons, 1):
        row_cells = table3.rows[i].cells
        row_cells[0].text = wrong
        row_cells[1].text = right
    
    doc.add_heading('4.3 数据呈现规范', level=2)
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = 'Table Grid'
    hdr_cells = table4.rows[0].cells
    headers = ['数据类型', '呈现方式', '示例']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], '2E75B6')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    data_types = [
        ('市场规模', '三情景表格', '保守/中性/乐观 + 数字 + 关键假设'),
        ('公司对比', '矩阵图或表格', '横轴能力，纵轴规模，气泡大小=估值'),
        ('时间线', '甘特图或阶段图', '2025/2026/2027关键节点'),
        ('产业链', '分层图', '上游→中游→下游，中美玩家分色标注'),
    ]
    for i, row_data in enumerate(data_types, 1):
        row_cells = table4.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_page_break()
    
    # 五、可视化规范
    doc.add_heading('五、可视化规范', level=1)
    doc.add_heading('5.1 必含图表（每份报告至少3个）', level=2)
    charts = [
        ('AMORA五维雷达图', '放在封面，直观展示行业在各维度的成熟度'),
        ('产业链生态位图', 'M章核心，分层展示中美玩家分布'),
        ('资本效率对比图', 'A章核心，融资额 vs 出货量/收入散点图'),
    ]
    for title, desc in charts:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f'：{desc}')
    
    doc.add_heading('5.2 可选图表', level=2)
    optional_charts = [
        '技术成熟度曲线（Hype Cycle风格）',
        '市场规模预测三情景图',
        '应用场景矩阵（横轴成熟度，纵轴市场规模）',
        '中美对比四象限图',
    ]
    for chart in optional_charts:
        doc.add_paragraph(chart, style='List Bullet')
    
    # 六、灵活调整原则
    doc.add_heading('六、灵活调整原则', level=1)
    doc.add_heading('6.1 固化的（必须遵守）', level=2)
    fixed_rules = [
        '五维字母顺序：A-M-O-R-A',
        '每章必须有"核心发现"段',
        '必须有中美对比结论',
        '封面必须有AMORA雷达图',
    ]
    for rule in fixed_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_heading('6.2 灵活的（根据行业调整）', level=2)
    flexible_rules = [
        '子章节数量：2-5节，根据信息量调整',
        '数据呈现形式：表格/图表/文字，根据数据特点选择',
        '章节侧重：早期技术行业A章重，成熟行业R/A章重',
        '案例深度：复杂行业详细拆，简单行业简化',
    ]
    for rule in flexible_rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    doc.add_page_break()
    
    # 七、质量检查清单
    doc.add_heading('七、质量检查清单', level=1)
    doc.add_paragraph('提交前检查：')
    checklist = [
        '封面有AMORA雷达图',
        'Executive Summary在2页以内',
        '每章有"核心发现"段',
        '有明确的中美对比结论（不是并列）',
        '每个结论有数据支撑',
        '有"So What"（对读者意味着什么）',
        '方法论说明（数据来源、预测口径）',
        '风险与展望章节',
        '附录含数据来源和术语表',
    ]
    for item in checklist:
        doc.add_paragraph(item, style='List Bullet')
    
    # 八、示例
    doc.add_heading('八、示例：人形机器人报告核心发现', level=1)
    
    examples = [
        ('A章核心发现', '人形机器人处于"0.5→1.0"过渡期，当前瓶颈不是AI算法，而是触觉传感器和电池续航。中国在上游电机/减速器已突破，但在AI芯片和触觉传感器仍有2-3年代差。'),
        ('M章核心发现', '美国占上游AI/传感器，中国占中游制造/量产。供应链脆弱性评估：若中美脱钩，中国在高端AI芯片和触觉传感器环节将断供，美国在量产制造环节将受阻。'),
        ('O章核心发现', '商业化率<5%，客户结构揭示"伪商业化"真相：宇树73.6%客户是科研教育，Figure在汽车工厂试点但规模极小。真正工业替代客户极少，ROI尚不可算。'),
        ('R章核心发现', '预测差600倍的根因是口径混乱（整机市场vs全产业链vs全社会价值）。我们以整机出货量×均价为基准，给出三情景预测：保守$XXB/中性$XXB/乐观$XXB。'),
        ('A章核心发现', 'Figure融资$1.5B估值$39B出货<150台，宇树融资$150M出货5500+台，资本效率差异巨大。中国公司资本效率显著更高，但美国公司估值溢价反映AI技术预期。'),
    ]
    
    for title, content in examples:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p = doc.add_paragraph(content)
        p.runs[0].italic = True
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 九、版本记录
    doc.add_heading('九、版本记录', level=1)
    table5 = doc.add_table(rows=2, cols=3)
    table5.style = 'Table Grid'
    hdr_cells = table5.rows[0].cells
    headers = ['版本', '日期', '更新内容']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr_cells[i], '2E75B6')
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    row_cells = table5.rows[1].cells
    row_cells[0].text = 'v1.0'
    row_cells[1].text = '2026-03-28'
    row_cells[2].text = '初始版本，建立五维框架和结构规范'
    
    doc.add_paragraph()
    
    # 结尾
    p = doc.add_paragraph('AMORA Research Team')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    p = doc.add_paragraph('标准化研究，差异化洞察')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    # 保存文档
    doc.save('AMORA-Report-Template-v1.0.docx')
    print('Word文档已生成: AMORA-Report-Template-v1.0.docx')

if __name__ == '__main__':
    create_doc()
